"""Render the v0.5 specification Markdown to DOCX and PDF.

Checked in beside the artifacts it produces, so a later version can be
regenerated rather than reverse-engineered. The Markdown is the source; both
derived files are byte-reproducible from it and the pinned renderer environment.
The manifest records the source, outputs, and build inputs.

Deliberately narrow: this reads the subset of Markdown the specification
actually uses — setext-free ATX headings, paragraphs, `-` bullets, `1.` ordered
items, pipe tables, blockquotes, and inline `code`/`**bold**` — and nothing
else. A general Markdown engine would be a larger dependency and a larger
surface for the two renderings to disagree about.

    pip install -r spec-bundle/render_v0_5.requirements.txt
    python spec-bundle/render_v0_5.py
"""

from __future__ import annotations

import hashlib
import io
import json
import os
import re
import sys
import zipfile
from functools import partial

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib import colors

HERE = os.path.dirname(os.path.abspath(__file__))
STEM = "pattern_like_astrology_app_product_platform_spec_v0.5"
SOURCE = os.path.join(HERE, f"{STEM}.md")
DOCX = os.path.join(HERE, f"{STEM}.docx")
PDF = os.path.join(HERE, f"{STEM}.pdf")
MANIFEST = os.path.join(HERE, f"{STEM}_manifest.json")
RENDERER = os.path.abspath(__file__)
REQUIREMENTS = os.path.join(HERE, "render_v0_5.requirements.txt")

INK = RGBColor(0x17, 0x31, 0x2A)
INK_HEX = colors.HexColor("#17312a")
SOFT_HEX = colors.HexColor("#4e625b")
RULE_HEX = colors.HexColor("#cec7b8")
DOCX_ZIP_TIMESTAMP = (2000, 1, 1, 0, 0, 0)


# ---------------------------------------------------------------------------
# Block parsing
# ---------------------------------------------------------------------------


def parse_blocks(text: str):
    """Markdown to a flat list of (kind, payload) blocks."""
    lines = text.replace("\r\n", "\n").split("\n")
    blocks = []
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            blocks.append(("rule", None))
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            blocks.append((f"h{len(heading.group(1))}", heading.group(2).strip()))
            index += 1
            continue

        if stripped.startswith("> "):
            quote = []
            while index < len(lines) and lines[index].strip().startswith("> "):
                quote.append(lines[index].strip()[2:])
                index += 1
            blocks.append(("quote", " ".join(quote)))
            continue

        if stripped.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            index += 1
            blocks.append(("code", "\n".join(code)))
            continue

        if stripped.startswith("| "):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [c.strip() for c in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                index += 1
            blocks.append(("table", rows))
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet or ordered:
            kind = "ul" if bullet else "ol"
            items = []
            while index < len(lines):
                candidate = lines[index].strip()
                match = re.match(
                    r"^[-*]\s+(.*)$" if kind == "ul" else r"^\d+\.\s+(.*)$", candidate
                )
                if match:
                    items.append(match.group(1))
                    index += 1
                elif candidate and not re.match(r"^(#{1,6}\s|[-*]\s|\d+\.\s|\||>|```)", candidate):
                    # A wrapped continuation line belongs to the item above it.
                    items[-1] += " " + candidate
                    index += 1
                else:
                    break
            blocks.append((kind, items))
            continue

        para = []
        while index < len(lines) and lines[index].strip() and not re.match(
            r"^(#{1,6}\s|[-*]\s|\d+\.\s|\||>|```|---$)", lines[index].strip()
        ):
            para.append(lines[index].strip())
            index += 1
        blocks.append(("p", " ".join(para)))

    return blocks


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def inline_runs(text: str):
    """Split inline markup into (text, bold, mono) runs. Links render as their label."""
    out = []
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            out.append((piece[2:-2], True, False))
        elif piece.startswith("`") and piece.endswith("`"):
            out.append((piece[1:-1], False, True))
        elif piece.startswith("["):
            label = piece[1 : piece.index("]")]
            out.append((label, False, False))
        else:
            out.append((piece, False, False))
    return out


def plain(text: str) -> str:
    return "".join(run[0] for run in inline_runs(text))


def rl_markup(text: str) -> str:
    """Inline markup as reportlab's mini-HTML, with the raw text escaped first."""
    out = []
    for content, bold, mono in inline_runs(text):
        escaped = (
            content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )
        if bold:
            out.append(f"<b>{escaped}</b>")
        elif mono:
            out.append(f'<font face="Courier" size="9">{escaped}</font>')
        else:
            out.append(escaped)
    return "".join(out)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def normalize_docx_archive() -> None:
    """Remove ZIP timestamps and platform attributes from the DOCX package."""
    with zipfile.ZipFile(DOCX, "r") as source:
        entries = [(info.filename, source.read(info.filename)) for info in source.infolist()]

    output = io.BytesIO()
    with zipfile.ZipFile(
        output,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        for filename, data in entries:
            info = zipfile.ZipInfo(filename, date_time=DOCX_ZIP_TIMESTAMP)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 3
            info.external_attr = 0o100600 << 16
            target.writestr(info, data, compresslevel=9)

    io.open(DOCX, "wb").write(output.getvalue())


def write_docx(blocks) -> None:
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)

    for level, size, space_before in (("Heading 1", 22, 0), ("Heading 2", 15, 18), ("Heading 3", 12, 14)):
        style = document.styles[level]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(6)

    def add_runs(paragraph, text):
        for content, bold, mono in inline_runs(text):
            run = paragraph.add_run(content)
            run.bold = bold
            if mono:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)

    for kind, payload in blocks:
        if kind == "h1":
            document.add_paragraph(plain(payload), style="Heading 1")
        elif kind == "h2":
            document.add_paragraph(plain(payload), style="Heading 2")
        elif kind in ("h3", "h4", "h5", "h6"):
            document.add_paragraph(plain(payload), style="Heading 3")
        elif kind == "p":
            add_runs(document.add_paragraph(), payload)
        elif kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(24)
            add_runs(paragraph, payload)
            for run in paragraph.runs:
                run.italic = True
        elif kind == "ul":
            for item in payload:
                add_runs(document.add_paragraph(style="List Bullet"), item)
        elif kind == "ol":
            for item in payload:
                add_runs(document.add_paragraph(style="List Number"), item)
        elif kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(payload)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "table" and payload:
            table = document.add_table(rows=0, cols=len(payload[0]))
            table.style = "Table Grid"
            for row_index, row in enumerate(payload):
                cells = table.add_row().cells
                for cell_index, cell in enumerate(row):
                    if cell_index >= len(cells):
                        continue
                    cells[cell_index].text = plain(cell)
                    if row_index == 0:
                        for paragraph in cells[cell_index].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        elif kind == "rule":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("· · ·")
            run.font.size = Pt(9)

    document.save(DOCX)
    normalize_docx_archive()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def write_pdf(blocks) -> None:
    sheet = getSampleStyleSheet()
    body = ParagraphStyle(
        "SpecBody",
        parent=sheet["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14.5,
        textColor=INK_HEX,
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    h1 = ParagraphStyle("SpecH1", parent=body, fontName="Helvetica-Bold", fontSize=20, leading=24, spaceAfter=10)
    h2 = ParagraphStyle("SpecH2", parent=body, fontName="Helvetica-Bold", fontSize=14, leading=18, spaceBefore=16, spaceAfter=6)
    h3 = ParagraphStyle("SpecH3", parent=body, fontName="Helvetica-Bold", fontSize=11, leading=15, spaceBefore=12, spaceAfter=4)
    quote = ParagraphStyle("SpecQuote", parent=body, leftIndent=18, textColor=SOFT_HEX, fontName="Helvetica-Oblique")
    code = ParagraphStyle("SpecCode", parent=body, fontName="Courier", fontSize=8.5, leading=11, leftIndent=12, textColor=SOFT_HEX)

    story = []
    for kind, payload in blocks:
        if kind == "h1":
            story.append(Paragraph(rl_markup(payload), h1))
        elif kind == "h2":
            story.append(Paragraph(rl_markup(payload), h2))
        elif kind in ("h3", "h4", "h5", "h6"):
            story.append(Paragraph(rl_markup(payload), h3))
        elif kind == "p":
            story.append(Paragraph(rl_markup(payload), body))
        elif kind == "quote":
            story.append(Paragraph(rl_markup(payload), quote))
        elif kind in ("ul", "ol"):
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(rl_markup(item), body), leftIndent=18) for item in payload],
                    bulletType="bullet" if kind == "ul" else "1",
                    start=None if kind == "ul" else 1,
                    leftIndent=18,
                )
            )
            story.append(Spacer(1, 6))
        elif kind == "code":
            for line in payload.split("\n"):
                story.append(Paragraph(rl_markup(line) or "&nbsp;", code))
            story.append(Spacer(1, 6))
        elif kind == "table" and payload:
            data = [[Paragraph(rl_markup(cell), body) for cell in row] for row in payload]
            table = Table(data, hAlign="LEFT", colWidths=[6.5 * inch / len(payload[0])] * len(payload[0]))
            table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, RULE_HEX),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8e3d8")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 10))
        elif kind == "rule":
            story.append(Spacer(1, 10))

    def furniture(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(SOFT_HEX)
        canvas.drawString(0.9 * inch, 0.6 * inch, "Pattern/Like — Product, UX, Data, and Platform Design Specification v0.5")
        canvas.drawRightString(LETTER[0] - 0.9 * inch, 0.6 * inch, str(document.page))
        canvas.restoreState()

    document = SimpleDocTemplate(
        PDF,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title="Pattern/Like Product, UX, Data, and Platform Design Specification v0.5",
        author="Pattern/Like",
        invariant=1,
    )
    document.build(
        story,
        onFirstPage=furniture,
        onLaterPages=furniture,
        canvasmaker=partial(Canvas, invariant=1),
    )


# ---------------------------------------------------------------------------
# Manifest
# ---------------------------------------------------------------------------


def digest(path: str):
    with io.open(path, "rb") as handle:
        data = handle.read()
    # LF-normalized for text so a checkout on any platform verifies; binaries
    # are hashed as they are.
    if path.endswith((".md", ".json", ".py", ".txt")):
        data = data.replace(b"\r\n", b"\n")
    return hashlib.sha256(data).hexdigest(), len(data)


def write_manifest() -> None:
    files = {}
    for path in (SOURCE, DOCX, PDF):
        sha, size = digest(path)
        files[os.path.basename(path)] = {"sha256": sha, "bytes": size}
    build_inputs = {}
    for path in (RENDERER, REQUIREMENTS):
        sha, size = digest(path)
        build_inputs[os.path.basename(path)] = {"sha256": sha, "bytes": size}

    manifest = {
        "//": [
            "Text hashes are over LF-normalized bytes so a checkout on any platform verifies;",
            "the DOCX and PDF are hashed as written. Both derived files are produced from the",
            "Markdown by spec-bundle/render_v0_5.py - regenerate rather than edit them.",
        ],
        "version": "0.5",
        "date": "2026-08-11",
        "supersedes": "0.2",
        "supersedes_scope": "sections 3, 7 daily reading generation, 10, 11, 12, 14, 15, 16",
        "source": os.path.basename(SOURCE),
        "renderer": "spec-bundle/render_v0_5.py",
        "build_inputs": build_inputs,
        "files": files,
        "carried_forward_unchanged": {
            "pattern_like_astrology_app_product_platform_spec_v0.2.md": "sections this version does not restate remain in force as published",
            "pattern_like_astrology_app_data_source_registry_v0.2.yaml": "unchanged; the source registry is not amended by this version",
            "pattern_like_astrology_app_platform_topology_v0.2.yaml": "unchanged",
        },
    }
    io.open(MANIFEST, "w", encoding="utf-8", newline="\n").write(
        json.dumps(manifest, indent=2) + "\n"
    )


def main() -> int:
    blocks = parse_blocks(io.open(SOURCE, encoding="utf-8").read())
    write_docx(blocks)
    write_pdf(blocks)
    write_manifest()
    for path in (DOCX, PDF, MANIFEST):
        sha, size = digest(path)
        print(f"{os.path.basename(path):60} {size:>9} bytes  {sha[:16]}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
